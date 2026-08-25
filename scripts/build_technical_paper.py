#!/usr/bin/env python3
"""Build the canonical ROOT technical concept paper DOCX."""

from __future__ import annotations

import argparse
import datetime as dt
import os
import pathlib
import sys
import tempfile
import zipfile
from collections.abc import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Twips


ROOT = pathlib.Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from paper.root_technical_paper import (  # noqa: E402
    CLAIMS,
    CONTROLLED_DISCLOSURES,
    DOCUMENT,
    EVIDENCE_STATUSES,
    FORBIDDEN_COPY,
    PROTOCOLS,
    REFERENCES,
    REVISION_HISTORY,
    SECTIONS,
)


TOKENS = {
    "base_preset": "compact_reference_guide",
    "cover_structure": "editorial_cover",
    "page_mm": (210, 297),
    "margins_mm": (20, 20, 20, 20),
    "header_footer_mm": 10,
    "content_width_dxa": 9638,
    "font": "Arial",
    "body_pt": 10,
    "body_after_pt": 5,
    "body_line_spacing": 1.18,
    "title_pt": 28,
    "h1_pt": 16,
    "h2_pt": 13,
    "h3_pt": 11.5,
    "ink": "20201F",
    "muted": "676761",
    "coral": "D85F48",
    "line": "D7D4CD",
    "light_fill": "F2F1ED",
    "named_overrides": {
        "a4_engineering_page": {
            "base_page": "US Letter",
            "resolved_page_mm": (210, 297),
            "resolved_margins_mm": (20, 20, 20, 20),
        },
        "engineering_typeface": {
            "base_font": "Calibri",
            "resolved_font": "Arial",
        },
        "compact_body_leading": {
            "base_size_pt": 11,
            "resolved_size_pt": 10,
            "base_after_pt": 6,
            "resolved_after_pt": 5,
            "base_line_spacing": 1.25,
            "resolved_line_spacing": 1.18,
        },
        "engineering_h1_spacing": {
            "base_before_after_pt": (18, 10),
            "resolved_before_after_pt": (18, 8),
        },
        "engineering_h2_spacing": {
            "base_before_after_pt": (14, 7),
            "resolved_before_after_pt": (14, 6),
        },
        "compact_h3_scale": {
            "base_pt": 12,
            "resolved_pt": 11.5,
        },
        "a4_table_geometry": {
            "base_width_dxa": 9360,
            "resolved_width_dxa": 9638,
            "base_header_fill": "E8EEF5",
            "resolved_header_fill": "F2F1ED",
            "preserved_indent_dxa": 120,
            "preserved_cell_margins_dxa": {
                "top": 80,
                "bottom": 80,
                "start": 120,
                "end": 120,
            },
        },
        "compact_list_leading": {
            "base_text_indent_dxa": 540,
            "resolved_text_indent_dxa": 540,
            "base_hanging_dxa": 271,
            "resolved_hanging_dxa": 270,
            "base_after_pt": 4,
            "resolved_after_pt": 4,
            "base_line_spacing": 1.25,
            "resolved_line_spacing": 1.18,
        },
        "restrained_editorial_cover": {
            "alignment": "left",
            "full_bleed_image": False,
            "top_spacer_pt": 78,
        },
    },
    "styles": {
        "normal": {"before_pt": 0},
        "title": {"before_pt": 0, "after_pt": 10, "line_spacing": 1.0},
        "subtitle": {"pt": 13, "before_pt": 0, "after_pt": 18, "line_spacing": 1.1},
        "heading_1": {"before_pt": 18, "after_pt": 8, "line_spacing": 1.05},
        "heading_2": {"before_pt": 14, "after_pt": 6, "line_spacing": 1.05},
        "heading_3": {"before_pt": 10, "after_pt": 5, "line_spacing": 1.05},
        "caption": {"pt": 8.5, "before_pt": 6, "after_pt": 4, "line_spacing": 1.0},
        "cover_kicker": {"pt": 9, "before_pt": 0, "after_pt": 12},
        "lead": {"pt": 10.5, "before_pt": 0, "after_pt": 7},
        "table_text": {"pt": 8.1, "before_pt": 0, "after_pt": 1.5, "line_spacing": 1.05},
        "contents_entry": {"pt": 9.5, "before_pt": 0, "after_pt": 3, "line_spacing": 1.05},
    },
    "header_footer": {"font_pt": 8, "paragraph_before_pt": 0, "paragraph_after_pt": 0},
    "tables": {
        "indent_dxa": 120,
        "cell_margins_dxa": {"top": 80, "bottom": 80, "start": 120, "end": 120},
        "border_size_eighth_pt": 4,
        "border_space_pt": 0,
        "after_paragraph_pt": 2,
    },
    "lists": {
        "tab_position_dxa": 540,
        "left_indent_dxa": 540,
        "hanging_dxa": 270,
        "after_pt": 4,
    },
    "contents": {"right_tab_mm": 170},
    "figures": {
        "product_width_mm": 105,
        "control_loop_weights": (1, 1, 1, 1),
        "alt_text": (
            "Illustrative cutaway of the proposed ROOT room-comfort controller; "
            "not component or production evidence."
        ),
    },
    "cover": {
        "top_spacer_after_pt": 78,
        "statement_before_pt": 18,
        "statement_after_pt": 56,
        "statement_pt": 11,
        "metadata_before_pt": 0,
        "metadata_after_pt": 4,
        "metadata_pt": 10,
        "detail_before_pt": 0,
        "detail_after_pt": 3,
        "detail_pt": 9,
        "owner_before_pt": 0,
        "owner_pt": 9,
    },
}

OUTPUT_PATH = ROOT / "output" / "docx" / "ROOT-Technical-Concept-Paper-D0.1.docx"
PRODUCT_IMAGE_PATH = ROOT / "assets" / "root-matte-technical-cutaway-v1.png"

DISPLAY_HEADINGS = {
    "abstract": "Abstract",
    "room-level-problem": "1. The room-level problem",
    "architecture": "2. System architecture",
    "connectivity": "3. Connectivity and local-control boundary",
    "sensing": "4. Sensing and infrared interface",
    "placement": "5. Placement and installation boundary",
    "claim-register": "6. Claim and evidence register",
    "evaluation": "7. Evaluation methodology",
    "limitations": "8. Limitations, privacy, and safety boundary",
    "company": "9. Company and document position",
    "references": "References",
    "revision-history": "Revision history",
    "appendix-a": "Appendix A. Controlled disclosures",
}

EVIDENCE_DEFINITIONS = {
    "Cited background": "A statement supported by an identified external source.",
    "Implemented prototype behavior": (
        "A status reserved for behavior demonstrated by a revision-linked prototype record; "
        "no claim in D0.1 is assigned this status."
    ),
    "Hypothesis / design target": "An intended behavior or bounded design thesis that remains unmeasured.",
    "Planned evaluation": "A claim framed as a protocol and acceptance decision to be completed later.",
}


def _rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def _ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_run_font(
    run,
    *,
    name: str = TOKENS["font"],
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    """Set a run font in both python-docx and explicit OOXML font slots."""

    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{slot}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = _rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_font(style, size: float, color: str, *, bold: bool = False, italic: bool = False) -> None:
    style.font.name = TOKENS["font"]
    style.font.size = Pt(size)
    style.font.color.rgb = _rgb(color)
    style.font.bold = bold
    style.font.italic = italic
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{slot}"), TOKENS["font"])


def configure_styles(document: Document) -> None:
    """Apply the A4 engineering-report style-sheet overrides."""

    styles = document.styles
    style_tokens = TOKENS["styles"]

    normal = styles["Normal"]
    _set_style_font(normal, TOKENS["body_pt"], TOKENS["ink"])
    normal.paragraph_format.space_before = Pt(style_tokens["normal"]["before_pt"])
    normal.paragraph_format.space_after = Pt(TOKENS["body_after_pt"])
    normal.paragraph_format.line_spacing = TOKENS["body_line_spacing"]
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    _set_style_font(title, TOKENS["title_pt"], TOKENS["ink"], bold=True)
    title.paragraph_format.space_before = Pt(style_tokens["title"]["before_pt"])
    title.paragraph_format.space_after = Pt(style_tokens["title"]["after_pt"])
    title.paragraph_format.line_spacing = style_tokens["title"]["line_spacing"]
    title.paragraph_format.keep_with_next = True
    title_p_pr = title._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    subtitle = styles["Subtitle"]
    _set_style_font(subtitle, style_tokens["subtitle"]["pt"], TOKENS["muted"])
    subtitle.paragraph_format.space_before = Pt(style_tokens["subtitle"]["before_pt"])
    subtitle.paragraph_format.space_after = Pt(style_tokens["subtitle"]["after_pt"])
    subtitle.paragraph_format.line_spacing = style_tokens["subtitle"]["line_spacing"]
    subtitle.paragraph_format.keep_with_next = True

    heading_tokens = (
        ("Heading 1", style_tokens["heading_1"], TOKENS["h1_pt"], TOKENS["coral"]),
        ("Heading 2", style_tokens["heading_2"], TOKENS["h2_pt"], TOKENS["ink"]),
        ("Heading 3", style_tokens["heading_3"], TOKENS["h3_pt"], TOKENS["muted"]),
    )
    for style_name, resolved, size, color in heading_tokens:
        style = styles[style_name]
        _set_style_font(style, size, color, bold=True)
        style.paragraph_format.space_before = Pt(resolved["before_pt"])
        style.paragraph_format.space_after = Pt(resolved["after_pt"])
        style.paragraph_format.line_spacing = resolved["line_spacing"]
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True
        p_pr = style._element.get_or_add_pPr()
        contextual_spacing = p_pr.find(qn("w:contextualSpacing"))
        if contextual_spacing is not None:
            p_pr.remove(contextual_spacing)

    caption = styles["Caption"]
    _set_style_font(caption, style_tokens["caption"]["pt"], TOKENS["muted"], bold=True)
    caption.paragraph_format.space_before = Pt(style_tokens["caption"]["before_pt"])
    caption.paragraph_format.space_after = Pt(style_tokens["caption"]["after_pt"])
    caption.paragraph_format.line_spacing = style_tokens["caption"]["line_spacing"]
    caption.paragraph_format.keep_with_next = True

    if "Cover Kicker" not in styles:
        styles.add_style("Cover Kicker", WD_STYLE_TYPE.PARAGRAPH)
    kicker = styles["Cover Kicker"]
    _set_style_font(kicker, style_tokens["cover_kicker"]["pt"], TOKENS["coral"], bold=True)
    kicker.paragraph_format.space_before = Pt(style_tokens["cover_kicker"]["before_pt"])
    kicker.paragraph_format.space_after = Pt(style_tokens["cover_kicker"]["after_pt"])
    kicker.paragraph_format.keep_with_next = True

    if "Lead" not in styles:
        styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
    lead = styles["Lead"]
    _set_style_font(lead, style_tokens["lead"]["pt"], TOKENS["muted"], italic=True)
    lead.paragraph_format.space_before = Pt(style_tokens["lead"]["before_pt"])
    lead.paragraph_format.space_after = Pt(style_tokens["lead"]["after_pt"])
    lead.paragraph_format.line_spacing = TOKENS["body_line_spacing"]
    lead.paragraph_format.keep_with_next = True

    if "Table Text" not in styles:
        styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    table_text = styles["Table Text"]
    _set_style_font(table_text, style_tokens["table_text"]["pt"], TOKENS["ink"])
    table_text.paragraph_format.space_before = Pt(style_tokens["table_text"]["before_pt"])
    table_text.paragraph_format.space_after = Pt(style_tokens["table_text"]["after_pt"])
    table_text.paragraph_format.line_spacing = style_tokens["table_text"]["line_spacing"]
    table_text.paragraph_format.widow_control = True

    if "Contents Entry" not in styles:
        styles.add_style("Contents Entry", WD_STYLE_TYPE.PARAGRAPH)
    toc = styles["Contents Entry"]
    _set_style_font(toc, style_tokens["contents_entry"]["pt"], TOKENS["ink"])
    toc.paragraph_format.space_before = Pt(style_tokens["contents_entry"]["before_pt"])
    toc.paragraph_format.space_after = Pt(style_tokens["contents_entry"]["after_pt"])
    toc.paragraph_format.line_spacing = style_tokens["contents_entry"]["line_spacing"]


def add_field(paragraph, field_code: str, placeholder: str = "1") -> None:
    """Append a simple Word field with stable cached display text."""

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {field_code} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    set_run_font(run, size=TOKENS["header_footer"]["font_pt"], color=TOKENS["muted"])
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(separate)
    cached = paragraph.add_run(placeholder)
    set_run_font(cached, size=TOKENS["header_footer"]["font_pt"], color=TOKENS["muted"])
    end_run = paragraph.add_run()
    set_run_font(end_run, size=TOKENS["header_footer"]["font_pt"], color=TOKENS["muted"])
    end_run._r.append(end)


def _populate_header(header) -> None:
    header_p = header.paragraphs[0]
    header_p.clear()
    header_p.paragraph_format.space_after = Pt(TOKENS["header_footer"]["paragraph_after_pt"])
    run = header_p.add_run(f"ROOT | Room-level AC comfort | {DOCUMENT['revision']}")
    set_run_font(run, size=TOKENS["header_footer"]["font_pt"], color=TOKENS["muted"], bold=True)


def _populate_footer(footer) -> None:
    footer_p = footer.paragraphs[0]
    footer_p.clear()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(TOKENS["header_footer"]["paragraph_before_pt"])
    prefix = footer_p.add_run(
        f"{DOCUMENT['document_id']}  |  {DOCUMENT['publication_status']}  |  Page "
    )
    set_run_font(prefix, size=TOKENS["header_footer"]["font_pt"], color=TOKENS["muted"])
    add_field(footer_p, "PAGE")
    tail = footer_p.add_run(" of ")
    set_run_font(tail, size=TOKENS["header_footer"]["font_pt"], color=TOKENS["muted"])
    add_field(footer_p, "NUMPAGES")


def _set_page_furniture(section) -> None:
    section.header.is_linked_to_previous = False
    _populate_header(section.header)
    section.footer.is_linked_to_previous = False
    _populate_footer(section.footer)


def configure_section(section) -> None:
    """Set deterministic A4 portrait geometry and running page furniture."""

    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(TOKENS["page_mm"][0])
    section.page_height = Mm(TOKENS["page_mm"][1])
    top, right, bottom, left = TOKENS["margins_mm"]
    section.top_margin = Mm(top)
    section.right_margin = Mm(right)
    section.bottom_margin = Mm(bottom)
    section.left_margin = Mm(left)
    section.header_distance = Mm(TOKENS["header_footer_mm"])
    section.footer_distance = Mm(TOKENS["header_footer_mm"])
    _set_page_furniture(section)


def set_cell_margins(cell, **margins_dxa: int) -> None:
    """Set explicit Word cell margins."""

    resolved = dict(TOKENS["tables"]["cell_margins_dxa"])
    resolved.update({key: int(value) for key, value in margins_dxa.items()})
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = _ensure_child(tc_pr, "w:tcMar")
    for side in ("top", "bottom", "start", "end"):
        node = _ensure_child(tc_mar, f"w:{side}")
        node.set(qn("w:w"), str(resolved[side]))
        node.set(qn("w:type"), "dxa")


def _set_dxa_width(parent, tag: str, width: int) -> None:
    node = _ensure_child(parent, tag)
    node.set(qn("w:type"), "dxa")
    node.set(qn("w:w"), str(int(width)))


def set_fixed_table_geometry(table, column_widths_dxa: Sequence[int]) -> None:
    """Synchronize tblW, tblInd, tblGrid, and every tcW in DXA."""

    widths = [int(width) for width in column_widths_dxa]
    if not widths or any(width <= 0 for width in widths):
        raise ValueError("table column widths must be positive")
    if sum(widths) != TOKENS["content_width_dxa"]:
        raise ValueError("table column widths must sum to the A4 content width")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    _set_dxa_width(tbl_pr, "w:tblW", sum(widths))
    indent = _ensure_child(tbl_pr, "w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(TOKENS["tables"]["indent_dxa"]))
    layout = _ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    borders = _ensure_child(tbl_pr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = _ensure_child(borders, f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(TOKENS["tables"]["border_size_eighth_pt"]))
        node.set(qn("w:space"), str(TOKENS["tables"]["border_space_pt"]))
        node.set(qn("w:color"), TOKENS["line"])

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for index, width in enumerate(widths):
        table.columns[index].width = Twips(width)
    for row in table.rows:
        row.height = None
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_dxa_width(cell._tc.get_or_add_tcPr(), "w:tcW", widths[index])
            set_cell_margins(cell)


def _allocate_widths(weights: Sequence[float]) -> list[int]:
    total = TOKENS["content_width_dxa"]
    weight_total = float(sum(weights))
    widths = [int(round(total * weight / weight_total)) for weight in weights]
    widths[-1] += total - sum(widths)
    return widths


def _mark_repeating_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def _shade_cell(cell, fill: str) -> None:
    shading = _ensure_child(cell._tc.get_or_add_tcPr(), "w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def _format_table(table) -> None:
    _mark_repeating_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            if row_index == 0:
                _shade_cell(cell, TOKENS["light_fill"])
            for paragraph in cell.paragraphs:
                paragraph.style = "Table Text"
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=TOKENS["styles"]["table_text"]["pt"],
                        color=TOKENS["ink"],
                        bold=row_index == 0,
                    )


def _add_table(
    document: Document,
    caption: str,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    weights: Sequence[float],
) -> None:
    caption_p = document.add_paragraph(caption, style="Caption")
    caption_p.paragraph_format.keep_with_next = True
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            row.cells[index].text = str(value)
    set_fixed_table_geometry(table, _allocate_widths(weights))
    _format_table(table)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(TOKENS["tables"]["after_paragraph_pt"])


def add_control_loop_figure(document: Document) -> None:
    """Add the conceptual four-stage control-loop figure as a technical grid."""

    stages = (
        ("1 | Sense", "Sample temperature, humidity, and bounded presence context."),
        ("2 | Interpret", "Estimate local room context without inferring safety-critical state."),
        ("3 | Decide", "Apply bounded setpoints, command rates, override, and fallback rules."),
        ("4 | Transmit", "Emit a learned IR command without assuming confirmed AC state."),
    )
    table = document.add_table(rows=2, cols=len(stages))
    for index, (stage, description) in enumerate(stages):
        table.rows[0].cells[index].text = stage
        table.rows[1].cells[index].text = description
    set_fixed_table_geometry(table, _allocate_widths(TOKENS["figures"]["control_loop_weights"]))
    _format_table(table)

    caption = document.add_paragraph(
        "Figure 1. Conceptual four-stage control loop. Conceptual architecture, not a measured result. "
        "Status: Hypothesis / design target. Claim ID: HT-01.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    document.add_paragraph(
        "Signal path: environmental and bounded presence inputs -> local interpretation -> bounded local "
        "decision -> learned IR transmission. The proposed loop does not provide confirmed AC-state feedback."
    )


def add_product_figure(document: Document, image_path: pathlib.Path) -> None:
    """Add the approved illustrative product cutaway as a bounded inline image."""

    image_path = pathlib.Path(image_path)
    if not image_path.is_file():
        raise FileNotFoundError(f"product figure not found: {image_path}")
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True
    inline_shape = image_paragraph.add_run().add_picture(
        str(image_path),
        width=Mm(TOKENS["figures"]["product_width_mm"]),
    )
    inline_shape._inline.docPr.set("descr", TOKENS["figures"]["alt_text"])
    inline_shape._inline.docPr.set("title", "Illustrative ROOT product cutaway")

    caption = document.add_paragraph(
        "Figure 2. Illustrative ROOT product cutaway. Illustrative image, not component or production "
        "evidence. Status: Hypothesis / design target. Claim ID: HT-03.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = False


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    resolved = TOKENS["styles"][f"heading_{level}"]
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(resolved["before_pt"])
    paragraph.paragraph_format.space_after = Pt(resolved["after_pt"])


def _add_section_copy(document: Document, section: dict, display_heading: str) -> None:
    _add_heading(document, display_heading)
    lead = document.add_paragraph(section["title"], style="Lead")
    lead.paragraph_format.keep_with_next = True
    for text in section["paragraphs"]:
        document.add_paragraph(text)


def _create_numbering(document: Document) -> dict[str, int]:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    next_abstract = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1
    result: dict[str, int] = {}

    for kind, num_format, level_text in (
        ("bullet", "bullet", "•"),
        ("decimal", "decimal", "%1."),
    ):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(next_abstract))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_format)
        level.append(fmt)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), level_text)
        level.append(text)
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        level.append(justification)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(TOKENS["lists"]["tab_position_dxa"]))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(TOKENS["lists"]["left_indent_dxa"]))
        ind.set(qn("w:hanging"), str(TOKENS["lists"]["hanging_dxa"]))
        p_pr.append(ind)
        level.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), TOKENS["font"])
        r_fonts.set(qn("w:hAnsi"), TOKENS["font"])
        r_pr.append(r_fonts)
        level.append(r_pr)
        abstract.append(level)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_num))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(next_abstract))
        num.append(abstract_ref)
        numbering.append(num)
        result[kind] = next_num
        next_abstract += 1
        next_num += 1
    return result


def _add_list_item(document: Document, text: str, num_id: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(TOKENS["lists"]["after_pt"])
    paragraph.paragraph_format.line_spacing = TOKENS["body_line_spacing"]
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = _ensure_child(p_pr, "w:numPr")
    ilvl = _ensure_child(num_pr, "w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = _ensure_child(num_pr, "w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    run = paragraph.add_run(text)
    set_run_font(run, size=TOKENS["body_pt"], color=TOKENS["ink"])


def add_cover(document: Document) -> None:
    """Add a restrained, left-aligned editorial cover."""

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(TOKENS["cover"]["top_spacer_after_pt"])
    kicker = document.add_paragraph("TECHNICAL CONCEPT PAPER", style="Cover Kicker")
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title = document.add_paragraph(DOCUMENT["title"], style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = document.add_paragraph(DOCUMENT["subtitle"], style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    statement = document.add_paragraph()
    statement.paragraph_format.space_before = Pt(TOKENS["cover"]["statement_before_pt"])
    statement.paragraph_format.space_after = Pt(TOKENS["cover"]["statement_after_pt"])
    run = statement.add_run(
        "A developer-preview architecture and evaluation plan for bounded local room-comfort control."
    )
    set_run_font(run, size=TOKENS["cover"]["statement_pt"], color=TOKENS["ink"])

    metadata = document.add_paragraph()
    metadata.paragraph_format.space_before = Pt(TOKENS["cover"]["metadata_before_pt"])
    metadata.paragraph_format.space_after = Pt(TOKENS["cover"]["metadata_after_pt"])
    first = metadata.add_run(f"{DOCUMENT['document_id']}  |  {DOCUMENT['revision']}")
    set_run_font(first, size=TOKENS["cover"]["metadata_pt"], color=TOKENS["ink"], bold=True)
    detail = document.add_paragraph()
    detail.paragraph_format.space_before = Pt(TOKENS["cover"]["detail_before_pt"])
    detail.paragraph_format.space_after = Pt(TOKENS["cover"]["detail_after_pt"])
    detail_run = detail.add_run(
        f"{DOCUMENT['publication_status']}  |  Issued {DOCUMENT['issue_date']}"
    )
    set_run_font(detail_run, size=TOKENS["cover"]["detail_pt"], color=TOKENS["muted"])
    owner = document.add_paragraph()
    owner.paragraph_format.space_before = Pt(TOKENS["cover"]["owner_before_pt"])
    owner_run = owner.add_run(DOCUMENT["owner"])
    set_run_font(owner_run, size=TOKENS["cover"]["owner_pt"], color=TOKENS["muted"])
    document.add_page_break()


def _toc_page(toc_page_map: dict[str, int] | None, key: str, heading: str) -> int | None:
    if toc_page_map is None:
        return None
    candidates = (key, heading, heading.split(".", 1)[0])
    for candidate in candidates:
        if candidate in toc_page_map:
            return int(toc_page_map[candidate])
    return None


def add_front_matter(
    document: Document,
    toc_page_map: dict[str, int] | None = None,
) -> None:
    """Add document control, contents, and nomenclature front matter."""

    _add_heading(document, "Document control")
    control_rows = (
        ("Document ID", DOCUMENT["document_id"]),
        ("Revision", DOCUMENT["revision"]),
        ("Publication status", DOCUMENT["publication_status"]),
        ("Owner", DOCUMENT["owner"]),
        ("Technical reviewer", DOCUMENT["technical_reviewer"]),
        ("Issue date", DOCUMENT["issue_date"]),
        ("Last reviewed", DOCUMENT["last_reviewed"]),
        ("Evidence cutoff", DOCUMENT["evidence_cutoff"]),
        ("Hardware revision", DOCUMENT["hardware_revision"]),
        ("Firmware revision", DOCUMENT["firmware_revision"]),
    )
    _add_table(document, "Table 1. Document-control register.", ("Field", "Controlled value"), control_rows, (1.6, 4.9))

    document.add_page_break()
    _add_heading(document, "Contents")
    for key in (
        "abstract",
        "room-level-problem",
        "architecture",
        "connectivity",
        "sensing",
        "placement",
        "claim-register",
        "evaluation",
        "limitations",
        "company",
        "references",
        "revision-history",
        "appendix-a",
    ):
        heading = DISPLAY_HEADINGS[key]
        paragraph = document.add_paragraph(style="Contents Entry")
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Mm(TOKENS["contents"]["right_tab_mm"]),
            WD_TAB_ALIGNMENT.RIGHT,
        )
        paragraph.add_run(heading)
        if key == "appendix-a" or heading[:1].isdigit():
            mapped_page = _toc_page(toc_page_map, key, heading)
            if mapped_page is not None:
                paragraph.add_run(f"\t{mapped_page}")

    _add_heading(document, "Nomenclature and evidence status")
    nomenclature_rows = [
        ("AC", "Air conditioner."),
        ("IR", "Infrared command learning or transmission."),
        ("RSSI", "Bluetooth received signal strength; a coarse and environment-dependent input."),
        ("ROOT", "The proposed room-comfort controller described by this paper."),
    ]
    nomenclature_rows.extend((status, EVIDENCE_DEFINITIONS[status]) for status in EVIDENCE_STATUSES)
    _add_table(
        document,
        "Table 2. Nomenclature and evidence-status taxonomy.",
        ("Term or status", "Meaning in this revision"),
        nomenclature_rows,
        (2.0, 4.5),
    )
    document.add_page_break()


def _claim_rows() -> list[tuple[str, ...]]:
    rows = []
    for claim in CLAIMS:
        control_record = (
            f"Wording: {claim['wording']}\n"
            f"Owner: {claim['owner']}\n"
            f"Review date: {claim['review_date']}\n"
            f"Superseded wording: {claim['superseded_wording'] or 'None recorded'}\n"
            f"Hardware revision: {claim['hardware_revision']}\n"
            f"Firmware revision: {claim['firmware_revision']}"
        )
        rows.append(
            (
                claim["id"],
                claim["status"],
                claim["scope"],
                claim["revision"],
                claim["evidence_id"],
                claim["evidence_date"],
                control_record,
            )
        )
    return rows


def _protocol_rows() -> list[tuple[str, ...]]:
    rows = []
    for protocol in PROTOCOLS:
        basis = f"{protocol['conditions']}\nComparator / ground truth: {protocol['comparator_ground_truth']}"
        outcomes = "; ".join(protocol["primary_outcomes"])
        rows.append(
            (
                protocol["claim_id"],
                protocol["title"],
                basis,
                outcomes,
                "Pre-registered before testing",
            )
        )
    return rows


def add_claim_register(document: Document) -> None:
    """Add every canonical claim and its complete evidence-control record."""

    _add_heading(document, DISPLAY_HEADINGS["claim-register"])
    document.add_paragraph(
        "The register below separates cited background, unmeasured design targets, and planned evaluations. "
        "It contains no claim assigned to Implemented prototype behavior."
    )
    _add_table(
        document,
        "Table 3. Claim and evidence register.",
        ("Claim ID", "Status", "Scope", "Revision", "Evidence ID", "Evidence date", "Control record"),
        _claim_rows(),
        (0.55, 1.0, 1.0, 0.85, 0.95, 0.75, 2.05),
    )


def add_protocol_matrix(document: Document) -> None:
    """Add the seven-row planned-evaluation matrix without unapproved thresholds."""

    document.add_paragraph(
        "Planned evaluation IDs: PE-01, PE-02, PE-03, PE-04, PE-05, PE-06, and PE-07. "
        "Each acceptance rule will be pre-registered before testing; no threshold is approved in D0.1."
    )
    _add_table(
        document,
        "Table 4. Planned evaluation matrix.",
        ("ID", "Evaluation", "Conditions and comparator", "Primary outputs", "Acceptance rule"),
        _protocol_rows(),
        (0.6, 1.0, 2.15, 1.55, 1.2),
    )


def add_technical_body(document: Document, numbering: dict[str, int]) -> None:
    """Add canonical technical sections, claim register, and evaluation protocols."""

    by_id = {section["id"]: section for section in SECTIONS}
    _add_section_copy(document, by_id["abstract"], DISPLAY_HEADINGS["abstract"])
    _add_section_copy(document, by_id["room-level-problem"], DISPLAY_HEADINGS["room-level-problem"])
    _add_section_copy(document, by_id["architecture"], DISPLAY_HEADINGS["architecture"])
    add_control_loop_figure(document)
    _add_heading(document, "Proposed control sequence", 2)
    for item in (
        "Sense temperature, humidity, and a bounded room-presence context near the occupied location.",
        "Estimate context without treating coarse proximity as a safety-relevant or irreversible signal.",
        "Decide locally within bounded setpoints, bounded command rates, manual override, and neutral fallback targets.",
        "Transmit a learned IR adjustment without treating the emitted command as confirmed AC state.",
    ):
        _add_list_item(document, item, numbering["decimal"])

    _add_section_copy(document, by_id["connectivity"], DISPLAY_HEADINGS["connectivity"])
    _add_section_copy(document, by_id["sensing"], DISPLAY_HEADINGS["sensing"])
    add_product_figure(document, PRODUCT_IMAGE_PATH)
    _add_section_copy(document, by_id["placement"], DISPLAY_HEADINGS["placement"])

    add_claim_register(document)

    _add_section_copy(document, by_id["evaluation"], DISPLAY_HEADINGS["evaluation"])
    _add_heading(document, "Protocol-wide reporting rules", 2)
    for item in (
        "Identify prototype hardware and firmware revisions, conditions, comparator, ground truth, sample interval, and repeated trials.",
        "Fix primary and secondary outcomes, acceptance criteria, exclusions, missing-data handling, and uncertainty before testing.",
        "Retain evidence artifacts and report trial counts, negative results, protocol deviations, data location, and analysis version.",
        "Keep simulator outputs distinct from measured results.",
    ):
        _add_list_item(document, item, numbering["bullet"])
    add_protocol_matrix(document)

    _add_section_copy(document, by_id["limitations"], DISPLAY_HEADINGS["limitations"])
    _add_section_copy(document, by_id["company"], DISPLAY_HEADINGS["company"])


def _reference_rows() -> list[tuple[str, ...]]:
    rows = []
    for reference in REFERENCES:
        creator = ", ".join(reference.get("authors", ())) or reference.get("organization", "")
        publication = f"{reference['publisher']} ({reference['publication_date']})"
        access = f"{reference['url']}\nAccessed {reference['access_date']}"
        rows.append((reference["id"], creator, reference["title"], publication, access))
    return rows


def add_back_matter(document: Document) -> None:
    """Add references, revision history, and controlled-disclosure appendix."""

    by_id = {section["id"]: section for section in SECTIONS}
    _add_section_copy(document, by_id["references"], DISPLAY_HEADINGS["references"])
    _add_table(
        document,
        "Table 5. Numbered reference register.",
        ("ID", "Author or organization", "Title", "Publication", "Locator and access"),
        _reference_rows(),
        (0.65, 1.65, 1.85, 1.0, 1.35),
    )

    _add_section_copy(document, by_id["revision-history"], DISPLAY_HEADINGS["revision-history"])
    revision_rows = [
        (
            item["revision"],
            item["issue_date"],
            item["owner"],
            item["reviewer"],
            item["evidence_cutoff"],
            item["summary"],
        )
        for item in REVISION_HISTORY
    ]
    _add_table(
        document,
        "Table 6. Revision history.",
        ("Revision", "Issue date", "Owner", "Reviewer", "Evidence cutoff", "Summary"),
        revision_rows,
        (0.7, 0.95, 1.3, 1.0, 1.0, 1.55),
    )

    document.add_page_break()
    _add_heading(document, DISPLAY_HEADINGS["appendix-a"])
    document.add_paragraph(
        "These controlled disclosures preserve the product's intended-use, safety, company, privacy, and omission boundaries."
    )
    disclosure_rows = [
        (
            item["id"],
            item["disclosure_type"],
            item["exact_wording"],
            item["source_or_approval_record"],
            item["owner"],
            f"Approved: {item['approval_date']}\nNext review: {item['next_review_date']}",
        )
        for item in CONTROLLED_DISCLOSURES
    ]
    _add_table(
        document,
        "Table 7. Controlled-disclosure register.",
        ("ID", "Type", "Exact wording", "Source / approval", "Owner", "Review control"),
        disclosure_rows,
        (0.55, 1.05, 2.1, 1.1, 0.85, 0.85),
    )


def _model_strings(value) -> list[str]:
    if isinstance(value, str):
        return [value]
    if isinstance(value, dict):
        return [text for nested in value.values() for text in _model_strings(nested)]
    if isinstance(value, (tuple, list)):
        return [text for nested in value for text in _model_strings(nested)]
    return []


def _validate_model() -> None:
    """Fail closed if the canonical content crosses a controlled-copy boundary."""

    corpus = "\n".join(
        _model_strings(
            (
                DOCUMENT,
                EVIDENCE_STATUSES,
                SECTIONS,
                CLAIMS,
                CONTROLLED_DISCLOSURES,
                REFERENCES,
                PROTOCOLS,
                REVISION_HISTORY,
            )
        )
    ).lower()
    for phrase in FORBIDDEN_COPY:
        if phrase.lower() in corpus:
            raise ValueError(f"forbidden copy in canonical model: {phrase}")
    if "Implemented prototype behavior" in {claim["status"] for claim in CLAIMS}:
        raise ValueError("D0.1 must not contain an implemented-prototype claim")


def _set_document_properties(document: Document) -> None:
    properties = document.core_properties
    timestamp = dt.datetime(2026, 8, 25, tzinfo=dt.timezone.utc)
    properties.title = DOCUMENT["title"]
    properties.subject = DOCUMENT["subtitle"]
    properties.author = DOCUMENT["owner"]
    properties.last_modified_by = DOCUMENT["owner"]
    properties.keywords = "ROOT, room comfort, split air conditioner, infrared control"
    properties.created = timestamp
    properties.modified = timestamp
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _normalize_docx_archive(source_path: pathlib.Path, normalized_path: pathlib.Path) -> None:
    """Repack with a fixed ZIP timestamp so repeated builds are byte-stable."""

    fixed_time = (2026, 8, 25, 0, 0, 0)
    with zipfile.ZipFile(source_path, "r") as source, zipfile.ZipFile(
        normalized_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as target:
        for name in source.namelist():
            original = source.getinfo(name)
            info = zipfile.ZipInfo(name, fixed_time)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = original.external_attr
            info.internal_attr = original.internal_attr
            info.create_system = original.create_system
            target.writestr(info, source.read(name))


def build_docx(
    output_path: pathlib.Path,
    toc_page_map: dict[str, int] | None = None,
) -> pathlib.Path:
    """Build the canonical DOCX and atomically replace ``output_path``."""

    _validate_model()
    output_path = pathlib.Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.settings.odd_and_even_pages_header_footer = False
    configure_styles(document)
    for section in document.sections:
        configure_section(section)
    _set_document_properties(document)
    numbering = _create_numbering(document)

    add_cover(document)
    add_front_matter(document, toc_page_map)
    add_technical_body(document, numbering)
    add_back_matter(document)

    raw_fd, raw_name = tempfile.mkstemp(prefix=f".{output_path.stem}-", suffix=".raw.docx", dir=output_path.parent)
    os.close(raw_fd)
    normalized_fd, normalized_name = tempfile.mkstemp(
        prefix=f".{output_path.stem}-", suffix=".normalized.docx", dir=output_path.parent
    )
    os.close(normalized_fd)
    raw_path = pathlib.Path(raw_name)
    normalized_path = pathlib.Path(normalized_name)
    try:
        document.save(raw_path)
        _normalize_docx_archive(raw_path, normalized_path)
        os.replace(normalized_path, output_path)
    finally:
        for temporary in (raw_path, normalized_path):
            if temporary.exists():
                temporary.unlink()
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx-only", action="store_true", help="Build only the canonical DOCX.")
    parser.add_argument("--output", type=pathlib.Path, default=OUTPUT_PATH)
    args = parser.parse_args()
    result = build_docx(args.output)
    print(result)


if __name__ == "__main__":
    main()
