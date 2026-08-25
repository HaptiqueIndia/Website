import importlib.util
import pathlib
import subprocess
import sys
import tempfile
import unittest
import zipfile
import xml.etree.ElementTree as ET
from unittest import mock

import pdfplumber
from docx import Document
from pypdf import PdfReader


ROOT = pathlib.Path(__file__).resolve().parents[1]
MODULE_PATH = ROOT / "paper" / "root_technical_paper.py"
BUILDER_PATH = ROOT / "scripts" / "build_technical_paper.py"
DOCX_PATH = ROOT / "output" / "docx" / "ROOT-Technical-Concept-Paper-D0.1.docx"
PDF_PATH = ROOT / "output" / "pdf" / "ROOT-Technical-Concept-Paper-D0.1.pdf"
APPROVED_FORBIDDEN_COPY = (
    "acceleration award", "affiliated with panasonic", "patent pending",
    "breathing sense", "zero outages", "perfect coverage",
    "universal compatibility", "fully-tested", "100% local",
)


class ContentModelTests(unittest.TestCase):
    def load_module(self):
        spec = importlib.util.spec_from_file_location("root_technical_paper", MODULE_PATH)
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        return module

    def test_document_control_is_complete(self):
        paper = self.load_module()
        self.assertEqual(paper.DOCUMENT["document_id"], "ROOT-TCP-001")
        self.assertEqual(paper.DOCUMENT["revision"], "D0.1")
        self.assertEqual(paper.DOCUMENT["issue_date"], "25 August 2026")
        self.assertEqual(paper.DOCUMENT["evidence_cutoff"], "25 August 2026")
        self.assertEqual(paper.DOCUMENT["publication_status"], "Developer preview")
        self.assertEqual(paper.DOCUMENT["technical_reviewer"], "Not yet assigned")

    def test_required_sections_are_ordered(self):
        paper = self.load_module()
        self.assertEqual(
            [section["id"] for section in paper.SECTIONS],
            [
                "abstract", "room-level-problem", "architecture", "connectivity",
                "sensing", "placement", "evaluation", "limitations", "company",
                "references", "revision-history",
            ],
        )

    def test_abstract_meets_technical_paper_word_budget(self):
        paper = self.load_module()
        abstract = next(section for section in paper.SECTIONS if section["id"] == "abstract")
        word_count = len(" ".join(abstract["paragraphs"]).split())

        self.assertGreaterEqual(word_count, 150)
        self.assertLessEqual(word_count, 220)

    def test_claim_register_has_required_fields_and_uses_no_prototype_status(self):
        paper = self.load_module()
        required = {
            "id", "status", "scope", "revision", "evidence_id", "evidence_date",
            "wording", "owner", "review_date", "superseded_wording",
            "hardware_revision", "firmware_revision",
        }
        self.assertTrue({"CB-01", "HT-01", "HT-02", "PE-01", "PE-07"}.issubset(
            {claim["id"] for claim in paper.CLAIMS}
        ))
        required_nonempty = required - {"superseded_wording"}
        for claim in paper.CLAIMS:
            self.assertTrue(required.issubset(claim))
            for field in required_nonempty:
                with self.subTest(claim_id=claim["id"], field=field):
                    self.assertTrue(claim[field])
            self.assertTrue(claim["superseded_wording"] is None or claim["superseded_wording"].strip())
        self.assertNotIn("Implemented prototype behavior", {claim["status"] for claim in paper.CLAIMS})
        for claim in paper.CLAIMS:
            expected_revision = "Not applicable" if claim["id"] == "CB-01" else "Target revision not yet assigned"
            self.assertEqual(claim["hardware_revision"], expected_revision)
            self.assertEqual(claim["firmware_revision"], expected_revision)

    def test_controlled_disclosures_have_approval_provenance(self):
        paper = self.load_module()
        required = {
            "id", "disclosure_type", "exact_wording", "source_or_approval_record",
            "owner", "approval_date", "next_review_date",
        }
        disclosure_types = {disclosure["disclosure_type"] for disclosure in paper.CONTROLLED_DISCLOSURES}
        self.assertEqual(
            disclosure_types,
            {"Intended use and medical boundary", "Safety role", "Company identity", "Privacy and data boundary", "Recognition and patent omission"},
        )
        for disclosure in paper.CONTROLLED_DISCLOSURES:
            self.assertTrue(required.issubset(disclosure))

    def test_protocols_have_complete_planning_contract(self):
        paper = self.load_module()
        required = {
            "claim_id", "title",
            "conditions", "comparator_ground_truth", "sample_interval", "repeated_trials",
            "primary_outcomes", "secondary_outcomes", "acceptance_criterion",
            "exclusions_missing_data", "uncertainty", "retained_evidence_artifact",
            "reporting_requirements", "hardware_revision", "firmware_revision",
        }
        for protocol in paper.PROTOCOLS:
            self.assertTrue(required.issubset(protocol))
            for field in required:
                with self.subTest(protocol_id=protocol["claim_id"], field=field):
                    self.assertTrue(protocol[field])
            self.assertEqual(protocol["acceptance_criterion"], "Defined before testing")
            self.assertEqual(protocol["hardware_revision"], "Target revision not yet assigned")
            self.assertEqual(protocol["firmware_revision"], "Target revision not yet assigned")

    def test_references_include_required_authorship_or_organization(self):
        paper = self.load_module()
        references = {reference["id"]: reference for reference in paper.REFERENCES}
        self.assertEqual(
            references["REF-01"]["authors"],
            ("Haiyan Yan", "Yawei Li", "Thomas Parkinson", "Stefano Schiavon", "Hui Zhang", "Rui Sun", "Shengkai Zhao", "Wei Zhao", "Zhen Sun", "Fangning Shi"),
        )
        self.assertEqual(references["REF-02"]["organization"], "Bluetooth SIG")

    def test_copy_respects_disclosure_boundary(self):
        paper = self.load_module()
        self.assertEqual(paper.FORBIDDEN_COPY, APPROVED_FORBIDDEN_COPY)
        def strings(value):
            if isinstance(value, str):
                return [value]
            if isinstance(value, dict):
                return [item for nested in value.values() for item in strings(nested)]
            if isinstance(value, (tuple, list)):
                return [item for nested in value for item in strings(nested)]
            return []

        corpus = "\n".join(strings((
            paper.EVIDENCE_STATUSES, paper.DOCUMENT, paper.SECTIONS, paper.CLAIMS, paper.REFERENCES,
            paper.PROTOCOLS, paper.CONTROLLED_DISCLOSURES, paper.REVISION_HISTORY,
        ))).lower()
        for phrase in APPROVED_FORBIDDEN_COPY:
            self.assertNotIn(phrase.lower(), corpus)


class DocxArtifactTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.document = Document(DOCX_PATH)

    def test_a4_geometry(self):
        section = self.document.sections[0]
        self.assertAlmostEqual(section.page_width.mm, 210, delta=0.2)
        self.assertAlmostEqual(section.page_height.mm, 297, delta=0.2)
        for margin in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin):
            self.assertAlmostEqual(margin.mm, 20, delta=0.2)

    def test_heading_structure_and_metadata(self):
        text = "\n".join(p.text for p in self.document.paragraphs)
        for required in (
            "ROOT: A local architecture for room-level AC comfort",
            "ROOT-TCP-001", "D0.1", "Abstract", "1. The room-level problem",
            "2. System architecture", "7. Evaluation methodology",
            "8. Limitations, privacy, and safety boundary", "References", "Revision history",
        ):
            self.assertIn(required, text)
        self.assertEqual(sum(p.style.name == "Title" for p in self.document.paragraphs), 1)

    def test_cover_visibly_carries_required_document_control(self):
        with zipfile.ZipFile(DOCX_PATH) as archive:
            root = ET.fromstring(archive.read("word/document.xml"))
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        body = root.find("w:body", namespace)
        cover_fragments = []
        for child in body:
            cover_fragments.extend(node.text or "" for node in child.findall(".//w:t", namespace))
            if child.find(".//w:br[@w:type='page']", namespace) is not None:
                break
        cover_text = " ".join(cover_fragments)

        for required in (
            "Owner: Haptique Electronics Pvt. Ltd.",
            "Technical reviewer: Not yet assigned",
            "Issued: 25 August 2026",
            "Evidence cutoff: 25 August 2026",
        ):
            with self.subTest(required=required):
                self.assertIn(required, cover_text)

    def test_tables_have_repeating_header_rows(self):
        self.assertGreaterEqual(len(self.document.tables), 4)
        for table in self.document.tables:
            first_row_xml = table.rows[0]._tr.xml
            self.assertIn("tblHeader", first_row_xml)

    def test_figures_and_protocols_are_complete(self):
        text = "\n".join(p.text for p in self.document.paragraphs)
        for marker in (
            "Figure 1. Conceptual four-stage control loop",
            "Figure 2. Illustrative ROOT product cutaway",
            "Conceptual architecture, not a measured result",
            "Illustrative image, not component or production evidence",
            "PE-01", "PE-02", "PE-03", "PE-04", "PE-05", "PE-06", "PE-07",
        ):
            self.assertIn(marker, text)

    def test_claim_register_includes_evidence_fields(self):
        expected_header = [
            "Claim ID", "Status", "Scope", "Revision", "Evidence ID", "Evidence date", "Control record",
        ]
        register = next(
            table for table in self.document.tables
            if [cell.text for cell in table.rows[0].cells] == expected_header
        )
        expected_claim_ids = {"CB-01", "HT-01", "HT-02", "PE-01", "PE-07"}
        rows = {row.cells[0].text: [cell.text for cell in row.cells] for row in register.rows[1:]}
        self.assertTrue(expected_claim_ids.issubset(rows))
        for claim_id, cells in rows.items():
            with self.subTest(claim_id=claim_id):
                self.assertEqual(len(cells), len(expected_header))
                self.assertTrue(all(cells[index].strip() for index in range(1, 7)))
                self.assertIn("Hardware revision:", cells[6])
                self.assertIn("Firmware revision:", cells[6])

    def test_product_figure_is_inline_bounded_and_has_alt_text(self):
        self.assertEqual(len(self.document.inline_shapes), 1)
        self.assertLessEqual(self.document.inline_shapes[0].width.mm, 120.0)
        caption = next(
            paragraph for paragraph in self.document.paragraphs
            if paragraph.text.startswith("Figure 2. Illustrative ROOT product cutaway")
        )
        self.assertIs(caption.paragraph_format.keep_with_next, False)
        with zipfile.ZipFile(DOCX_PATH) as archive:
            root = ET.fromstring(archive.read("word/document.xml"))
        namespace = {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}
        descriptions = [node.get("descr", "") for node in root.findall(".//wp:docPr", namespace)]
        self.assertIn(
            "Illustrative cutaway of the proposed ROOT room-comfort controller; not component or production evidence.",
            descriptions,
        )

    def test_each_published_evaluation_exposes_the_complete_protocol_contract(self):
        expected_header = [
            "Evaluation", "Controlled setup", "Measurement and decision record",
        ]
        matrix = next(
            table for table in self.document.tables
            if [cell.text for cell in table.rows[0].cells] == expected_header
        )
        self.assertEqual(len(matrix.rows), 8)
        self.assertEqual(
            [row.cells[0].text.splitlines()[0] for row in matrix.rows[1:]],
            [f"PE-{i:02d}" for i in range(1, 8)],
        )
        required_labels = (
            "Hardware revision:",
            "Firmware revision:",
            "Conditions:",
            "Comparator / ground truth:",
            "Sample interval:",
            "Repeated trials:",
            "Primary outcomes:",
            "Secondary outcomes:",
            "Acceptance criterion:",
            "Exclusions / missing data:",
            "Uncertainty:",
            "Retained artifact:",
        )
        for row in matrix.rows[1:]:
            evaluation_id = row.cells[0].text.splitlines()[0]
            published_record = "\n".join(cell.text for cell in row.cells)
            for label in required_labels:
                with self.subTest(evaluation_id=evaluation_id, label=label):
                    self.assertIn(label, published_record)


class DocxBuilderTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        spec = importlib.util.spec_from_file_location("build_technical_paper", BUILDER_PATH)
        cls.builder = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(cls.builder)

    def test_front_matter_accepts_explicit_toc_page_map(self):
        document = Document()
        self.builder.configure_styles(document)
        for section in document.sections:
            self.builder.configure_section(section)

        self.builder.add_front_matter(
            document,
            {"room-level-problem": 42, "appendix-a": 99},
        )

        contents = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("1. The room-level problem\t42", contents)
        self.assertIn("Appendix A. Controlled disclosures\t99", contents)

    def test_build_docx_creates_valid_a4_artifact(self):
        with tempfile.TemporaryDirectory() as temporary_directory:
            output_path = pathlib.Path(temporary_directory) / "built.docx"
            result = self.builder.build_docx(
                output_path,
                {"room-level-problem": 4, "appendix-a": 11},
            )

            self.assertEqual(result, output_path)
            self.assertTrue(output_path.is_file())
            document = Document(output_path)
            section = document.sections[0]
            self.assertAlmostEqual(section.page_width.mm, 210, delta=0.2)
            self.assertAlmostEqual(section.page_height.mm, 297, delta=0.2)
            for margin in (
                section.top_margin,
                section.right_margin,
                section.bottom_margin,
                section.left_margin,
            ):
                self.assertAlmostEqual(margin.mm, 20, delta=0.2)

            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            for required in (
                "ROOT: A local architecture for room-level AC comfort",
                "ROOT-TCP-001",
                "1. The room-level problem\t4",
                "Appendix A. Controlled disclosures\t11",
                "Implemented prototype behavior",
                "Revision history",
            ):
                self.assertIn(required, text)
            self.assertGreaterEqual(len(document.tables), 7)
            for table in document.tables:
                self.assertIn("tblHeader", table.rows[0]._tr.xml)

            with zipfile.ZipFile(output_path) as archive:
                footer_root = ET.fromstring(archive.read("word/footer1.xml"))
                settings_xml = archive.read("word/settings.xml").decode("utf-8")
            namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            instructions = [
                "".join(node.itertext()).strip()
                for node in footer_root.findall(".//w:instrText", namespace)
            ]
            self.assertIn("PAGE", instructions)
            self.assertIn("NUMPAGES", instructions)
            self.assertNotEqual(instructions.index("PAGE"), instructions.index("NUMPAGES"))
            self.assertIn("updateFields", settings_xml)

    def test_build_docx_is_byte_deterministic(self):
        with tempfile.TemporaryDirectory() as temporary_directory:
            directory = pathlib.Path(temporary_directory)
            first = self.builder.build_docx(directory / "first.docx")
            second = self.builder.build_docx(directory / "second.docx")

            self.assertEqual(first.read_bytes(), second.read_bytes())

    def test_pdf_only_preserves_existing_docx_bytes(self):
        with tempfile.TemporaryDirectory() as temporary_directory:
            directory = pathlib.Path(temporary_directory)
            docx_path = directory / "canonical.docx"
            pdf_path = directory / "canonical.pdf"
            render_dir = directory / "render"
            original_bytes = b"already-scrubbed-canonical-docx"
            docx_path.write_bytes(original_bytes)

            def fake_render(source, supplied_render_dir, destination):
                self.assertEqual(pathlib.Path(source), docx_path)
                self.assertEqual(pathlib.Path(supplied_render_dir), render_dir)
                pathlib.Path(destination).write_bytes(b"pdf")
                return pathlib.Path(destination)

            arguments = [
                str(BUILDER_PATH), "--pdf-only", "--output", str(docx_path),
                "--pdf-output", str(pdf_path), "--render-dir", str(render_dir),
            ]
            with (
                mock.patch.object(sys, "argv", arguments),
                mock.patch.object(self.builder, "build_docx", side_effect=AssertionError("must not rebuild")),
                mock.patch.object(self.builder, "render_canonical_pdf", side_effect=fake_render) as render,
            ):
                self.builder.main()

            render.assert_called_once()
            self.assertEqual(docx_path.read_bytes(), original_bytes)

    def test_pdf_render_cleanup_is_scoped_and_publish_is_atomic_on_failure(self):
        with tempfile.TemporaryDirectory() as temporary_directory:
            directory = pathlib.Path(temporary_directory)
            render_root = directory / "render-root"
            render_dir = render_root / "current"
            render_dir.mkdir(parents=True)
            (render_dir / "stale.png").write_bytes(b"stale")
            sibling = render_root / "preserve.txt"
            sibling.write_bytes(b"preserve")
            docx_path = directory / "canonical.docx"
            docx_path.write_bytes(b"docx")
            pdf_path = directory / "canonical.pdf"
            original_pdf = b"previous-good-pdf"
            pdf_path.write_bytes(original_pdf)

            def fake_run(*_args, **_kwargs):
                (render_dir / "page-1.png").write_bytes(b"png")
                (render_dir / f"{docx_path.stem}.pdf").write_bytes(b"replacement-pdf")
                return subprocess.CompletedProcess([], 0, "rendered", "")

            with (
                mock.patch.object(self.builder, "RENDER_ROOT", render_root),
                mock.patch.object(self.builder.subprocess, "run", side_effect=fake_run),
                mock.patch.object(self.builder, "_validate_pdf"),
                mock.patch.object(
                    self.builder.os,
                    "replace",
                    side_effect=OSError("simulated atomic publish failure"),
                ),
            ):
                with self.assertRaisesRegex(OSError, "simulated atomic publish failure"):
                    self.builder.render_canonical_pdf(docx_path, render_dir, pdf_path)

            self.assertFalse((render_dir / "stale.png").exists())
            self.assertEqual(sibling.read_bytes(), b"preserve")
            self.assertEqual(pdf_path.read_bytes(), original_pdf)


class PdfArtifactTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        spec = importlib.util.spec_from_file_location("build_technical_paper_pdf", BUILDER_PATH)
        cls.builder = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(cls.builder)
        cls.reader = PdfReader(PDF_PATH)
        with pdfplumber.open(PDF_PATH) as pdf:
            cls.text = "\n".join(page.extract_text() or "" for page in pdf.pages)

    def test_page_geometry_and_count(self):
        self.assertGreaterEqual(len(self.reader.pages), 12)
        self.assertLessEqual(len(self.reader.pages), 16)
        for page in self.reader.pages:
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
            self.assertAlmostEqual(width, 595.28, delta=1.0)
            self.assertAlmostEqual(height, 841.89, delta=1.0)

    def test_required_pdf_content(self):
        for marker in (
            "ROOT-TCP-001", "D0.1", "Developer preview", "Abstract",
            "System architecture", "Evaluation methodology",
            "Limitations, privacy, and safety boundary", "References",
            "Revision history", "CB-01", "HT-01", "PE-07",
        ):
            self.assertIn(marker, self.text)

    def test_pdf_cover_and_protocol_appendix_publish_required_controls(self):
        with pdfplumber.open(PDF_PATH) as pdf:
            cover_text = pdf.pages[0].extract_text() or ""
        for required in (
            "Owner: Haptique Electronics Pvt. Ltd.",
            "Technical reviewer: Not yet assigned",
            "Issued: 25 August 2026",
            "Evidence cutoff: 25 August 2026",
        ):
            with self.subTest(cover=required):
                self.assertIn(required, cover_text)

        for required in (
            "Hardware revision:",
            "Firmware revision:",
            "Sample interval:",
            "Repeated trials:",
            "Secondary outcomes:",
            "Exclusions / missing data:",
            "Uncertainty:",
            "Retained artifact:",
        ):
            with self.subTest(protocol_field=required):
                self.assertIn(required, self.text)

    def test_table_of_contents_has_page_numbers(self):
        self.assertRegex(
            self.text,
            r"Table of contents[\s\S]{0,5000}1\. The room-level problem[ .]+[0-9]+",
        )

    def test_pdf_has_no_forbidden_copy(self):
        lowered = self.text.lower()
        for phrase in (
            "acceleration award", "affiliated with panasonic", "patent pending",
            "breathing sense", "zero outages", "perfect coverage",
        ):
            self.assertNotIn(phrase, lowered)

    def test_heading_page_extraction_is_complete(self):
        expected_keys = {
            key for key, heading in self.builder.DISPLAY_HEADINGS.items()
            if key == "appendix-a" or heading[:1].isdigit()
        }
        self.assertEqual(set(self.builder.extract_heading_pages(PDF_PATH)), expected_keys)


if __name__ == "__main__":
    unittest.main()
